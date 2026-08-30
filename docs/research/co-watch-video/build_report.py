from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "report-source.md"
OUTPUT = ROOT / "EduSpace-CoWatch-Deep-Research-FA.docx"

NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B3A"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF7E0"
GOLD = "9A7411"
WHITE = "FFFFFF"
BORDER = "D0D5DD"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        old_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_rtl(paragraph, rtl: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1" if rtl else "0")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Arial", rtl=True) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if rtl:
        rtl_el = r_pr.find(qn("w:rtl"))
        if rtl_el is None:
            rtl_el = OxmlElement("w:rtl")
            r_pr.append(rtl_el)
        rtl_el.set(qn("w:val"), "1")
        lang = r_pr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            r_pr.append(lang)
        lang.set(qn("w:bidi"), "fa-IR")


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), "Arial")
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    r_pr.extend([fonts, color, underline, rtl])
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\)|`[^`]+`)")


def add_inline(paragraph, text: str, size: float = 11, color: str = INK) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(8.5, size - 0.5), color=DARK_BLUE, name="Consolas", rtl=False)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def style_paragraph(paragraph, after=6, before=0, line=1.10, rtl=True) -> None:
    set_rtl(paragraph, rtl)
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_together = False
    pf.widow_control = True


def add_text(doc, text: str, style=None, after=6, before=0, line=1.10, size=11, color=INK, bold=False):
    p = doc.add_paragraph(style=style)
    style_paragraph(p, after=after, before=before, line=line)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_callout(doc, text: str, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "bottom", "end"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "single")
    start.set(qn("w:sz"), "18")
    start.set(qn("w:color"), accent)
    borders.append(start)
    p = cell.paragraphs[0]
    style_paragraph(p, after=2, before=2, line=1.15)
    add_inline(p, text, size=11, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=9, color=MUTED, rtl=False)


def create_numbering(doc: Document, kind: str, start: int = 1) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start_el = OxmlElement("w:start")
    start_el.set(qn("w:val"), str(start))
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start_el, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def add_list_item(doc, text: str, num_id: int):
    p = doc.add_paragraph()
    style_paragraph(p, after=8, before=0, line=1.167)
    apply_numbering(p, num_id)
    add_inline(p, text)
    return p


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    style_tokens = {
        "Heading 1": (16, BLUE, 12, 6),
        "Heading 2": (13, BLUE, 10, 5),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in style_tokens.items():
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def setup_document(doc: Document) -> None:
    setup_styles(doc)
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    doc.core_properties.title = "تحقیق جامع قابلیت پخش همزمان ویدئو در EduSpace"
    doc.core_properties.subject = "تحلیل فنی، محصول، بازار و نقشه راه Co-Watch"
    doc.core_properties.author = "EduSpace Product & Engineering"
    doc.core_properties.keywords = "EduSpace, Co-Watch, HLS, CMAF, LiveKit, Video Upload"


def add_cover(doc: Document) -> None:
    add_text(doc, "EDUSPACE", after=0, before=2, line=1.0, size=12, color=GOLD, bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(90)
    set_rtl(spacer)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("تحقیق جامع پخش همزمان ویدئو")
    set_run_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Co-Watch و پخش پیش از پایان آپلود")
    set_run_font(r, size=17, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(46)
    r = p.add_run("تحلیل فنی، محصول، بازار، ریسک، هزینه و نقشه راه مخصوص EduSpace")
    set_run_font(r, size=12.5, color=MUTED)

    add_callout(
        doc,
        "نتیجه اصلی: قابلیت ارزشمند و شدنی است؛ اما باید با «مدت واقعاً آماده پخش»، کیفیت تطبیقی و fallback طراحی شود، نه با وعده عمومی «۵٪ فایل» یا «صفر لگ».",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(34)
    set_rtl(spacer)

    meta = [
        "مخاطب: تیم محصول، فنی و کسب‌وکار EduSpace",
        "نسخه: ۱.۰ | تاریخ: ۲۹ اوت ۲۰۲۶",
        "مبنای پروژه: شاخه codex/fix/camera-background-lifecycle",
        "وضعیت: داخلی - جهت تصمیم‌گیری",
    ]
    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line)
        set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def add_decision_map(doc: Document) -> None:
    p = doc.add_paragraph(style="Heading 1")
    set_rtl(p)
    add_inline(p, "نقشه تصمیم در یک صفحه", size=16, color=BLUE)
    add_callout(
        doc,
        "توصیه: Shared Player و VOD آماده را ابتدا عرضه کنید؛ Progressive Partial Playback را فقط پس از Spike موفق فعال کنید؛ برای شروع فوری، مسیر زنده موقت و handoff به HLS را در نظر بگیرید.",
    )

    headers = ["گزینه", "مزیت اصلی", "محدودیت اصلی", "جایگاه پیشنهادی"]
    rows = [
        ["Managed VOD", "سریع‌ترین راه به ABR/CDN", "پخش معمولاً پس از upload/processing", "Pilot و اثبات بازار"],
        ["Self-hosted VOD", "کنترل داده و vendor", "عملیات رسانه و ظرفیت‌سازی سنگین", "وقتی vendor مناسب نیست"],
        ["Progressive ingest", "پخش قبل از پایان upload", "برای هر MP4 عمومی نیست؛ seek محدود", "فاز آزمایشی پس از Spike"],
        ["Live + handoff", "شروع فوری از دستگاه میزبان", "تا handoff وابسته به uplink میزبان", "fallback تجربه کاربر"],
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, [1600, 2600, 2600, 2560])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        style_paragraph(p, after=1, before=1, line=1.05)
        r = p.add_run(text)
        set_run_font(r, size=9, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(values):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            style_paragraph(p, after=1, before=1, line=1.08)
            r = p.add_run(text)
            set_run_font(r, size=8.7, bold=(idx == 0), color=INK)

    p = doc.add_paragraph()
    style_paragraph(p, after=10, before=4, line=1.0)
    r = p.add_run("مبنای مقایسه: شواهد رسمی استانداردها، سرویس‌های ویدئو، رقبا و کد زنده EduSpace.")
    set_run_font(r, size=8.5, italic=True, color=MUTED)

    p = doc.add_paragraph(style="Heading 2")
    set_rtl(p)
    add_inline(p, "راهنمای مطالعه", size=13, color=BLUE)
    for text in [
        "بخش‌های ۱ تا ۳: مسئله کاربر، اعتبارسنجی گزارش قبلی و بازار.",
        "بخش‌های ۴ تا ۶: اصول فنی، گزینه‌ها و معماری مخصوص EduSpace.",
        "بخش‌های ۷ تا ۹: UX، امنیت، حقوق، هزینه و ظرفیت.",
        "بخش‌های ۱۰ تا ۱۴: نقشه راه، KPI، تست، تصمیم‌های باز و توصیه نهایی.",
    ]:
        add_list_item(doc, text, BULLET_NUM_ID)
    doc.add_page_break()


def add_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## پاسخ مدیریتی مستقیم"))
    lines = lines[start:]
    skipping_ledger = False
    in_code = False
    code_lines: list[str] = []
    list_kind = None
    decimal_num_id = None
    bullet_num_id = BULLET_NUM_ID

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()

        if stripped.startswith("## ضمیمه A -"):
            skipping_ledger = True
            i += 1
            continue
        if skipping_ledger and stripped.startswith("## ضمیمه B -"):
            skipping_ledger = False
        if skipping_ledger:
            i += 1
            continue

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph()
                style_paragraph(p, after=8, before=4, line=1.0, rtl=False)
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.right_indent = Inches(0.2)
                p_pr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F7F8FA")
                p_pr.append(shd)
                r = p.add_run("\n".join(code_lines))
                set_run_font(r, size=8.3, color=DARK_BLUE, name="Consolas", rtl=False)
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not stripped or stripped == "---":
            list_kind = None
            decimal_num_id = None
            i += 1
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if heading_match:
            hashes, title = heading_match.groups()
            level = min(3, len(hashes) - 1)
            p = doc.add_paragraph(style=f"Heading {level}")
            set_rtl(p)
            size = {1: 16, 2: 13, 3: 12}[level]
            color = {1: BLUE, 2: BLUE, 3: DARK_BLUE}[level]
            add_inline(p, title, size=size, color=color)
            list_kind = None
            decimal_num_id = None
            i += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            add_callout(doc, " ".join(quote_lines), fill=PALE_GOLD, accent=GOLD)
            list_kind = None
            decimal_num_id = None
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if bullet_match:
            add_list_item(doc, bullet_match.group(1), bullet_num_id)
            list_kind = "bullet"
            i += 1
            continue
        if number_match:
            if list_kind != "decimal" or decimal_num_id is None:
                decimal_num_id = create_numbering(doc, "decimal")
            add_list_item(doc, number_match.group(1), decimal_num_id)
            list_kind = "decimal"
            i += 1
            continue

        # Markdown table support, used only for genuinely tabular data.
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            parsed = [[c.strip() for c in line.strip("|").split("|")] for line in table_lines]
            parsed = [row for row in parsed if not all(re.fullmatch(r":?-{3,}:?", c or "") for c in row)]
            if parsed:
                cols = len(parsed[0])
                widths = [9360 // cols] * cols
                widths[-1] += 9360 - sum(widths)
                table = doc.add_table(rows=1, cols=cols)
                table.style = "Table Grid"
                table.autofit = False
                set_table_geometry(table, widths)
                for col, text in enumerate(parsed[0]):
                    cell = table.rows[0].cells[col]
                    set_cell_shading(cell, NAVY)
                    p = cell.paragraphs[0]
                    style_paragraph(p, after=1, before=1, line=1.0)
                    add_inline(p, text, size=8.2, color=WHITE)
                set_repeat_table_header(table.rows[0])
                for row in parsed[1:]:
                    cells = table.add_row().cells
                    for col, text in enumerate(row):
                        p = cells[col].paragraphs[0]
                        style_paragraph(p, after=1, before=1, line=1.05)
                        add_inline(p, text, size=8.2, color=INK)
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(4)
            continue

        p = doc.add_paragraph()
        style_paragraph(p)
        add_inline(p, stripped)
        list_kind = None
        decimal_num_id = None
        i += 1


doc = Document()
setup_document(doc)
BULLET_NUM_ID = create_numbering(doc, "bullet")
add_cover(doc)
add_decision_map(doc)
add_markdown_body(doc, SOURCE.read_text(encoding="utf-8"))

# Final preset-sensitive cleanup.
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

doc.save(OUTPUT)
print(OUTPUT)
